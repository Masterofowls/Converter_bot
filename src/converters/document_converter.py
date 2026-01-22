"""
Document Converter - handles PDF, DOCX, TXT, MD, JSON, XML, YAML, CSV
"""

import csv
import json
import logging
import re
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import pdfplumber
import yaml
from docx import Document
from fpdf import FPDF
from markdown import markdown

from .base import BaseConverter, ConversionResult

logger = logging.getLogger(__name__)


class DocumentConverter(BaseConverter):
    """Converter for document formats"""

    @property
    def supported_input_formats(self) -> set:
        return {
            "csv",
            "pdf",
            "docx",
            "xml",
            "json",
            "yaml",
            "yml",
            "md",
            "txt",
        }

    @property
    def supported_output_formats(self) -> set:
        return {
            "csv",
            "pdf",
            "docx",
            "xml",
            "json",
            "yaml",
            "md",
            "txt",
            "html",
        }

    async def convert(
        self,
        input_path: Path,
        output_format: str,
        options: Optional[Dict[str, Any]] = None,
    ) -> ConversionResult:
        """Convert document to specified format"""
        start_time = datetime.now()
        options = options or {}

        try:
            self.validate_input(input_path)
            input_format = input_path.suffix.lstrip(".").lower()
            output_format = output_format.lower()

            if input_format == "yml":
                input_format = "yaml"

            output_path = self.get_output_path(input_path, output_format)

            conversion_method = self._get_conversion_method(input_format, output_format)
            if not conversion_method:
                return ConversionResult(
                    success=False,
                    error_message=(f"Unsupported: {input_format} to {output_format}"),
                )

            await conversion_method(input_path, output_path, options)
            elapsed = (datetime.now() - start_time).total_seconds()

            return ConversionResult(
                success=True,
                input_path=input_path,
                output_path=output_path,
                input_format=input_format,
                output_format=output_format,
                conversion_time=elapsed,
                file_size=self.get_file_size(output_path),
            )

        except Exception as e:
            logger.error(f"Document conversion failed: {e}")
            return ConversionResult(
                success=False,
                input_path=input_path,
                input_format=input_path.suffix.lstrip("."),
                output_format=output_format,
                error_message=str(e),
            )

    def _get_conversion_method(self, input_fmt: str, output_fmt: str):
        """Get the appropriate conversion method"""
        methods = {
            ("txt", "pdf"): self._txt_to_pdf,
            ("txt", "docx"): self._txt_to_docx,
            ("txt", "md"): self._txt_to_md,
            ("md", "pdf"): self._md_to_pdf,
            ("md", "html"): self._md_to_html,
            ("md", "txt"): self._md_to_txt,
            ("md", "docx"): self._md_to_docx,
            ("docx", "pdf"): self._docx_to_pdf,
            ("docx", "txt"): self._docx_to_txt,
            ("docx", "md"): self._docx_to_md,
            ("pdf", "txt"): self._pdf_to_txt,
            ("pdf", "docx"): self._pdf_to_docx,
            ("pdf", "md"): self._pdf_to_md,
            ("json", "xml"): self._json_to_xml,
            ("json", "yaml"): self._json_to_yaml,
            ("json", "csv"): self._json_to_csv,
            ("json", "txt"): self._json_to_txt,
            ("json", "md"): self._json_to_md,
            ("xml", "json"): self._xml_to_json,
            ("xml", "yaml"): self._xml_to_yaml,
            ("xml", "csv"): self._xml_to_csv,
            ("xml", "txt"): self._xml_to_txt,
            ("yaml", "json"): self._yaml_to_json,
            ("yaml", "xml"): self._yaml_to_xml,
            ("yaml", "txt"): self._yaml_to_txt,
            ("csv", "json"): self._csv_to_json,
            ("csv", "xml"): self._csv_to_xml,
            ("csv", "yaml"): self._csv_to_yaml,
            ("csv", "txt"): self._csv_to_txt,
            ("csv", "md"): self._csv_to_md,
            ("csv", "pdf"): self._csv_to_pdf,
        }
        return methods.get((input_fmt, output_fmt))

    async def _txt_to_pdf(self, inp: Path, out: Path, opts: dict):
        text = inp.read_text(encoding="utf-8", errors="ignore")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        for line in text.split("\n"):
            pdf.multi_cell(0, 5, line)
        pdf.output(str(out))

    async def _txt_to_docx(self, inp: Path, out: Path, opts: dict):
        text = inp.read_text(encoding="utf-8", errors="ignore")
        doc = Document()
        for para in text.split("\n"):
            doc.add_paragraph(para)
        doc.save(str(out))

    async def _txt_to_md(self, inp: Path, out: Path, opts: dict):
        text = inp.read_text(encoding="utf-8", errors="ignore")
        out.write_text(text, encoding="utf-8")

    async def _md_to_pdf(self, inp: Path, out: Path, opts: dict):
        md_text = inp.read_text(encoding="utf-8", errors="ignore")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        for line in md_text.split("\n"):
            if line.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.multi_cell(0, 8, line[2:])
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 7, line[3:])
            elif line.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.multi_cell(0, 6, line[4:])
            elif line.startswith("```"):
                pdf.set_font("Courier", size=9)
            elif line.startswith("- ") or line.startswith("* "):
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 5, "  • " + line[2:])
            else:
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 5, line)
        pdf.output(str(out))

    async def _md_to_html(self, inp: Path, out: Path, opts: dict):
        md_text = inp.read_text(encoding="utf-8", errors="ignore")
        html_content = markdown(md_text, extensions=["tables", "fenced_code"])
        full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Document</title>
<style>
body {{ font-family: Arial; padding: 20px; margin: auto; }}
code {{ background: #f4f4f4; padding: 2px 5px; }}
pre {{ background: #f4f4f4; padding: 10px; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; }}
</style></head><body>{html_content}</body></html>"""
        out.write_text(full, encoding="utf-8")

    async def _md_to_txt(self, inp: Path, out: Path, opts: dict):
        text = inp.read_text(encoding="utf-8", errors="ignore")
        text = re.sub(r"#{1,6}\s+", "", text)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        out.write_text(text, encoding="utf-8")

    async def _md_to_docx(self, inp: Path, out: Path, opts: dict):
        md_text = inp.read_text(encoding="utf-8", errors="ignore")
        doc = Document()
        for line in md_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        doc.save(str(out))

    async def _docx_to_pdf(self, inp: Path, out: Path, opts: dict):
        doc = Document(str(inp))
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 7, para.text)
                pdf.set_font("Helvetica", size=10)
            else:
                pdf.multi_cell(0, 5, para.text)
        pdf.output(str(out))

    async def _docx_to_txt(self, inp: Path, out: Path, opts: dict):
        doc = Document(str(inp))
        text = "\n".join([p.text for p in doc.paragraphs])
        out.write_text(text, encoding="utf-8")

    async def _docx_to_md(self, inp: Path, out: Path, opts: dict):
        doc = Document(str(inp))
        lines = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                ch = para.style.name[-1]
                lvl = int(ch) if ch.isdigit() else 1
                lines.append("#" * lvl + " " + para.text)
            else:
                lines.append(para.text)
        out.write_text("\n\n".join(lines), encoding="utf-8")

    async def _pdf_to_txt(self, inp: Path, out: Path, opts: dict):
        parts = []
        with pdfplumber.open(str(inp)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        out.write_text("\n\n".join(parts), encoding="utf-8")

    async def _pdf_to_docx(self, inp: Path, out: Path, opts: dict):
        doc = Document()
        with pdfplumber.open(str(inp)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for para in text.split("\n"):
                        doc.add_paragraph(para)
        doc.save(str(out))

    async def _pdf_to_md(self, inp: Path, out: Path, opts: dict):
        parts = []
        with pdfplumber.open(str(inp)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    parts.append(f"## Page {i + 1}\n\n{text}")
        out.write_text("\n\n".join(parts), encoding="utf-8")

    async def _json_to_xml(self, inp: Path, out: Path, opts: dict):
        data = json.loads(inp.read_text(encoding="utf-8"))
        xml_str = self._dict_to_xml(data, "root")
        header = '<?xml version="1.0" encoding="UTF-8"?>\n'
        out.write_text(header + xml_str, encoding="utf-8")

    async def _json_to_yaml(self, inp: Path, out: Path, opts: dict):
        data = json.loads(inp.read_text(encoding="utf-8"))
        yml = yaml.dump(data, default_flow_style=False, allow_unicode=True)
        out.write_text(yml, encoding="utf-8")

    async def _json_to_csv(self, inp: Path, out: Path, opts: dict):
        data = json.loads(inp.read_text(encoding="utf-8"))
        if isinstance(data, list) and data:
            with open(out, "w", newline="", encoding="utf-8") as f:
                w = csv.DictWriter(f, fieldnames=data[0].keys())
                w.writeheader()
                w.writerows(data)
        else:
            with open(out, "w", newline="", encoding="utf-8") as f:
                w = csv.DictWriter(f, fieldnames=data.keys())
                w.writeheader()
                w.writerow(data)

    async def _json_to_txt(self, inp: Path, out: Path, opts: dict):
        data = json.loads(inp.read_text(encoding="utf-8"))
        out.write_text(json.dumps(data, indent=2), encoding="utf-8")

    async def _json_to_md(self, inp: Path, out: Path, opts: dict):
        data = json.loads(inp.read_text(encoding="utf-8"))
        md = "# JSON Data\n\n```json\n"
        md += json.dumps(data, indent=2, ensure_ascii=False)
        md += "\n```"
        out.write_text(md, encoding="utf-8")

    async def _xml_to_json(self, inp: Path, out: Path, opts: dict):
        tree = ET.parse(str(inp))
        data = self._xml_to_dict(tree.getroot())
        out.write_text(json.dumps(data, indent=2), encoding="utf-8")

    async def _xml_to_yaml(self, inp: Path, out: Path, opts: dict):
        tree = ET.parse(str(inp))
        data = self._xml_to_dict(tree.getroot())
        yml = yaml.dump(data, default_flow_style=False, allow_unicode=True)
        out.write_text(yml, encoding="utf-8")

    async def _xml_to_csv(self, inp: Path, out: Path, opts: dict):
        tree = ET.parse(str(inp))
        data = self._xml_to_dict(tree.getroot())
        flat = self._flatten_dict(data)
        with open(out, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["Key", "Value"])
            for k, v in flat.items():
                w.writerow([k, v])

    async def _xml_to_txt(self, inp: Path, out: Path, opts: dict):
        tree = ET.parse(str(inp))
        ET.indent(tree, space="  ")
        out.write_text(ET.tostring(tree.getroot(), encoding="unicode"))

    async def _yaml_to_json(self, inp: Path, out: Path, opts: dict):
        data = yaml.safe_load(inp.read_text(encoding="utf-8"))
        out.write_text(json.dumps(data, indent=2), encoding="utf-8")

    async def _yaml_to_xml(self, inp: Path, out: Path, opts: dict):
        data = yaml.safe_load(inp.read_text(encoding="utf-8"))
        xml_str = self._dict_to_xml(data, "root")
        header = '<?xml version="1.0" encoding="UTF-8"?>\n'
        out.write_text(header + xml_str, encoding="utf-8")

    async def _yaml_to_txt(self, inp: Path, out: Path, opts: dict):
        out.write_text(inp.read_text(encoding="utf-8"), encoding="utf-8")

    async def _csv_to_json(self, inp: Path, out: Path, opts: dict):
        with open(inp, "r", encoding="utf-8") as f:
            data = list(csv.DictReader(f))
        out.write_text(json.dumps(data, indent=2), encoding="utf-8")

    async def _csv_to_xml(self, inp: Path, out: Path, opts: dict):
        with open(inp, "r", encoding="utf-8") as f:
            data = list(csv.DictReader(f))
        xml_str = self._dict_to_xml({"rows": {"row": data}}, "data")
        header = '<?xml version="1.0" encoding="UTF-8"?>\n'
        out.write_text(header + xml_str, encoding="utf-8")

    async def _csv_to_yaml(self, inp: Path, out: Path, opts: dict):
        with open(inp, "r", encoding="utf-8") as f:
            data = list(csv.DictReader(f))
        yml = yaml.dump(data, default_flow_style=False, allow_unicode=True)
        out.write_text(yml, encoding="utf-8")

    async def _csv_to_txt(self, inp: Path, out: Path, opts: dict):
        out.write_text(inp.read_text(encoding="utf-8"), encoding="utf-8")

    async def _csv_to_md(self, inp: Path, out: Path, opts: dict):
        with open(inp, "r", encoding="utf-8") as f:
            rows = list(csv.reader(f))
        if not rows:
            out.write_text("", encoding="utf-8")
            return
        lines = []
        headers = rows[0]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        out.write_text("\n".join(lines), encoding="utf-8")

    async def _csv_to_pdf(self, inp: Path, out: Path, opts: dict):
        with open(inp, "r", encoding="utf-8") as f:
            rows = list(csv.reader(f))
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        if not rows:
            pdf.set_font("Helvetica", size=10)
            pdf.cell(0, 10, "Empty file")
            pdf.output(str(out))
            return
        col_count = len(rows[0])
        col_width = (pdf.w - 20) / max(col_count, 1)
        # Header
        pdf.set_fill_color(242, 242, 242)
        pdf.set_font("Helvetica", "B", 8)
        for cell in rows[0]:
            pdf.cell(col_width, 8, str(cell)[:20], border=1, fill=True)
        pdf.ln()
        # Data rows
        pdf.set_font("Helvetica", size=7)
        for row in rows[1:500]:  # Limit rows
            for cell in row:
                pdf.cell(col_width, 6, str(cell)[:20], border=1)
            pdf.ln()
        pdf.output(str(out))

    def _dict_to_xml(self, data: Any, root_name: str) -> str:
        """Convert dictionary to XML string"""

        def _to_xml(d, parent):
            if isinstance(d, dict):
                for key, val in d.items():
                    child = ET.SubElement(parent, str(key))
                    _to_xml(val, child)
            elif isinstance(d, list):
                for item in d:
                    child = ET.SubElement(parent, "item")
                    _to_xml(item, child)
            else:
                parent.text = str(d) if d is not None else ""

        root = ET.Element(root_name)
        _to_xml(data, root)
        ET.indent(root, space="  ")
        return ET.tostring(root, encoding="unicode")

    def _xml_to_dict(self, element: ET.Element) -> dict:
        """Convert XML element to dictionary"""
        result = {}
        for child in element:
            value = self._xml_to_dict(child) if len(child) else child.text
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(value)
            else:
                result[child.tag] = value
        return result if result else element.text

    def _flatten_dict(self, d: dict, parent: str = "", sep: str = ".") -> dict:
        """Flatten nested dictionary"""
        items = []
        for k, v in d.items():
            new_key = f"{parent}{sep}{k}" if parent else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key, sep).items())
            elif isinstance(v, list):
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        flat = self._flatten_dict(item, f"{new_key}[{i}]", sep)
                        items.extend(flat.items())
                    else:
                        items.append((f"{new_key}[{i}]", item))
            else:
                items.append((new_key, v))
        return dict(items)
